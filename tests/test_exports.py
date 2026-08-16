from io import BytesIO

import pytest
from docx import Document

from auto_report_agent.docx_export import _parse_inline_spans, markdown_to_docx_bytes
from auto_report_agent.pdf_export import _inline, markdown_to_pdf_bytes

SAMPLE_MARKDOWN = """# 测试报告

这是一段包含 **加粗**、`代码` 和 [链接](https://example.com) 的内容。

| 指标 | 数值 | 说明 |
| --- | --- | --- |
| 准确率 | **92%** | 表格内容应自动换行并正确渲染 |

- 要点一
- 要点二
"""

ALL_HEADING_LEVELS = "\n\n".join(f"{'#' * level} 标题 {level}" for level in range(1, 7))


def test_markdown_to_docx_bytes() -> None:
    data = markdown_to_docx_bytes(SAMPLE_MARKDOWN)
    assert data.startswith(b"PK")
    assert len(data) > 1000


def test_markdown_to_pdf_bytes() -> None:
    data = markdown_to_pdf_bytes(SAMPLE_MARKDOWN)
    assert data.startswith(b"%PDF")
    assert b"<b>" not in data
    assert len(data) > 1000


def test_docx_renders_every_heading_level() -> None:
    """Levels 4-6 used to fall through and print literal '####' in the body."""
    document = Document(BytesIO(markdown_to_docx_bytes(ALL_HEADING_LEVELS)))
    rendered = [p for p in document.paragraphs if p.text.strip()]

    assert [p.style.name for p in rendered] == [f"Heading {level}" for level in range(1, 7)]
    assert [p.text for p in rendered] == [f"标题 {level}" for level in range(1, 7)]
    assert not any("#" in p.text for p in rendered)


def test_pdf_renders_every_heading_level() -> None:
    data = markdown_to_pdf_bytes(ALL_HEADING_LEVELS)
    assert data.startswith(b"%PDF")


def test_heading_marker_requires_a_space() -> None:
    """'#标签' is not a heading and must survive as ordinary text."""
    document = Document(BytesIO(markdown_to_docx_bytes("#标签\n")))
    rendered = [p for p in document.paragraphs if p.text.strip()]

    assert [p.style.name for p in rendered] == ["Normal"]
    assert rendered[0].text == "#标签"


@pytest.mark.parametrize(
    ("source", "expected"),
    [
        ("**粗体**", "<b>粗体</b>"),
        ("*斜体*", "<i>斜体</i>"),
        ("***粗斜体***", "<b><i>粗斜体</i></b>"),
        ("普通文本", "普通文本"),
        ("**未闭合", "**未闭合"),
        ("a * b", "a * b"),
    ],
)
def test_pdf_inline_markup(source: str, expected: str) -> None:
    assert _inline(source) == expected


def test_pdf_export_survives_crossed_inline_markup() -> None:
    """Ambiguous emphasis can yield crossed tags, which ReportLab rejects.

    The exporter must drop the styling for that line instead of failing the
    whole download.
    """
    crossed = "**a *b** c*"
    assert _inline(crossed) == "<b>a <i>b</b> c</i>"  # genuinely malformed

    data = markdown_to_pdf_bytes(f"# 标题\n\n{crossed}\n")
    assert data.startswith(b"%PDF")


@pytest.mark.parametrize("text", ["**未闭合", "*未闭合", "a * b * c", "***x**", "<b>字面</b>"])
def test_pdf_export_survives_odd_markers(text: str) -> None:
    data = markdown_to_pdf_bytes(f"# 标题\n\n{text}\n")
    assert data.startswith(b"%PDF")


# --- DOCX inline parsing -----------------------------------------------------


def _spans(text: str) -> list[tuple[str, bool, bool, bool]]:
    return [(s.text, s.bold, s.italic, s.code) for s in _parse_inline_spans(text)]


def test_plain_text_is_a_single_span():
    assert _spans("just words") == [("just words", False, False, False)]


def test_bold_italic_and_code_are_marked():
    assert _spans("**b**") == [("b", True, False, False)]
    assert _spans("*i*") == [("i", False, True, False)]
    assert _spans("`c`") == [("c", False, False, True)]


def test_links_become_label_plus_url():
    assert _spans("[label](https://example.com)") == [
        ("label (https://example.com)", False, False, False)
    ]


def test_surrounding_text_is_preserved_around_markup():
    assert _spans("before **mid** after") == [
        ("before ", False, False, False),
        ("mid", True, False, False),
        (" after", False, False, False),
    ]


def test_the_earliest_marker_wins():
    """Scanning must not reorder spans by pattern priority."""
    assert _spans("`code` then **bold**") == [
        ("code", False, False, True),
        (" then ", False, False, False),
        ("bold", True, False, False),
    ]


def test_bold_is_not_mistaken_for_italic():
    assert _spans("**bold**") == [("bold", True, False, False)]


def test_unbalanced_markers_stay_literal():
    assert _spans("**unclosed") == [("**unclosed", False, False, False)]
    assert _spans("a * b") == [("a * b", False, False, False)]


def test_empty_input_yields_no_spans():
    assert _parse_inline_spans("") == []


def test_multiple_markers_of_each_kind():
    assert _spans("**a** and **b**") == [
        ("a", True, False, False),
        (" and ", False, False, False),
        ("b", True, False, False),
    ]


def test_docx_renders_inline_styles_into_runs():
    document = Document(BytesIO(markdown_to_docx_bytes("普通 **加粗** 与 `代码`")))
    runs = [r for p in document.paragraphs for r in p.runs if r.text.strip()]

    styled = {r.text: (r.bold, r.font.name) for r in runs}
    assert styled["加粗"][0] is True
    assert styled["代码"][1] == "Consolas"
