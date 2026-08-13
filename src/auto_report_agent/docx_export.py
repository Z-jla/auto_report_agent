from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO


@dataclass
class _InlineSpan:
    text: str
    bold: bool = False
    italic: bool = False
    code: bool = False


def markdown_to_docx_bytes(markdown_text: str, title: str = "Auto Report") -> bytes:
    """Convert a Markdown report to DOCX bytes.

    Supports headings, paragraphs, blockquotes, bullet/numbered lists,
    fenced code blocks and simple pipe tables.
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("缺少 Word 导出依赖 python-docx，请先运行：pip install -e .") from exc

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    _set_default_font(doc, qn, Pt)

    lines = markdown_text.replace("\r\n", "\n").split("\n")
    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                _add_code_block(doc, "\n".join(code_lines), Pt, RGBColor)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph("")
            i += 1
            continue

        if stripped.startswith("|") and _looks_like_table(lines, i):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            _add_table(doc, table_lines)
            continue

        if stripped.startswith("# "):
            _add_heading_with_inline(doc, stripped[2:], level=1)
        elif stripped.startswith("## "):
            _add_heading_with_inline(doc, stripped[3:], level=2)
        elif stripped.startswith("### "):
            _add_heading_with_inline(doc, stripped[4:], level=3)
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            _write_inline_runs(p, stripped[2:], italic_default=True)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.color.rgb = RGBColor(90, 90, 90)
        elif re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            _write_inline_runs(p, re.sub(r"^[-*]\s+", "", stripped))
        elif re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            _write_inline_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
        else:
            p = doc.add_paragraph()
            _write_inline_runs(p, stripped)
        i += 1

    if code_lines:
        _add_code_block(doc, "\n".join(code_lines), Pt, RGBColor)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _set_default_font(doc, qn, Pt) -> None:
    styles = doc.styles
    for style_name in [
        "Normal",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "List Bullet",
        "List Number",
    ]:
        if style_name not in styles:
            continue
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        if style_name == "Normal":
            style.font.size = Pt(10.5)


_INLINE_PATTERNS: tuple[tuple[str, re.Pattern[str]], ...] = (
    ("link", re.compile(r"\[([^\]]+)\]\(([^)]+)\)")),
    ("bold", re.compile(r"\*\*(.+?)\*\*")),
    ("italic", re.compile(r"(?<!\*)\*([^*\n]+?)\*(?!\*)")),
    ("code", re.compile(r"`([^`]+)`")),
)


def _parse_inline_spans(text: str) -> list[_InlineSpan]:
    """Parse common Markdown inline markers into styled spans."""
    spans: list[_InlineSpan] = []
    cursor = 0
    while cursor < len(text):
        next_match: tuple[str, re.Match[str]] | None = None
        for kind, pattern in _INLINE_PATTERNS:
            match = pattern.search(text, cursor)
            if match and (next_match is None or match.start() < next_match[1].start()):
                next_match = (kind, match)
        if next_match is None:
            spans.append(_InlineSpan(text=text[cursor:]))
            break

        kind, match = next_match
        if match.start() > cursor:
            spans.append(_InlineSpan(text=text[cursor : match.start()]))

        if kind == "link":
            label, url = match.group(1), match.group(2)
            spans.append(_InlineSpan(text=f"{label} ({url})"))
        elif kind == "bold":
            spans.append(_InlineSpan(text=match.group(1), bold=True))
        elif kind == "italic":
            spans.append(_InlineSpan(text=match.group(1), italic=True))
        elif kind == "code":
            spans.append(_InlineSpan(text=match.group(1), code=True))

        cursor = match.end()

    return [s for s in spans if s.text]


def _write_inline_runs(paragraph, text: str, *, italic_default: bool = False) -> None:
    """Render inline-formatted Markdown into the given paragraph as styled runs."""
    spans = _parse_inline_spans(text.strip())
    if not spans:
        return
    for span in spans:
        run = paragraph.add_run(span.text)
        if span.bold:
            run.bold = True
        if span.italic or italic_default:
            run.italic = True
        if span.code:
            run.font.name = "Consolas"


def _add_heading_with_inline(doc, text: str, *, level: int) -> None:
    heading = doc.add_heading("", level=level)
    _write_inline_runs(heading, text)


def _looks_like_table(lines: list[str], index: int) -> bool:
    return (
        index + 1 < len(lines)
        and lines[index].strip().startswith("|")
        and re.match(r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$", lines[index + 1].strip())
        is not None
    )


def _add_table(doc, table_lines: list[str]) -> None:
    rows: list[list[str]] = []
    for idx, line in enumerate(table_lines):
        if idx == 1:
            continue
        rows.append([cell.strip() for cell in line.strip("|").split("|")])

    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            _write_inline_runs(paragraph, cell_text)
            if r_idx == 0:
                for run in paragraph.runs:
                    run.bold = True
    doc.add_paragraph("")


def _add_code_block(doc, code_text: str, Pt, RGBColor) -> None:
    p = doc.add_paragraph()
    from docx.shared import Inches

    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(60, 60, 60)
