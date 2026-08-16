from __future__ import annotations

import re
import zipfile
from collections.abc import Iterable, Sequence
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Protocol

from docx import Document

SUPPORTED_DOCUMENT_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
DEFAULT_MAX_DOCS = 5
# Matches the PAPER_MAX_CHARS_PER_DOC default both entrypoints pass. Text beyond
# this is sampled from the head, middle and tail rather than cut off.
DEFAULT_MAX_CHARS_PER_DOC = 120000
DEFAULT_PREVIEW_CHARS = 1200
DEFAULT_MAX_FILE_BYTES = 25 * 1024 * 1024
DEFAULT_MAX_TOTAL_BYTES = 50 * 1024 * 1024
DEFAULT_MAX_PDF_PAGES = 200
DEFAULT_MAX_DOCX_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
DEFAULT_MAX_DOCX_COMPRESSION_RATIO = 100.0


class UploadedFile(Protocol):
    """The part of Streamlit's UploadedFile this module actually uses."""

    name: str

    def getvalue(self) -> bytes: ...


@dataclass
class ParsedDocument:
    name: str
    extension: str
    content: str
    original_chars: int
    extracted_chars: int
    truncated: bool


def _normalize_text(text: str) -> str:
    text = text.replace("\x00", "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _decode_text_bytes(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "gbk"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _get_pdf_reader_class():
    try:
        from pypdf import PdfReader

        return PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader

            return PdfReader
        except ImportError as exc:
            raise RuntimeError(
                "解析 PDF 需要安装 pypdf 或 PyPDF2，请先执行 `pip install pypdf`。"
            ) from exc


def _extract_pdf_text(data: bytes, *, max_pdf_pages: int) -> str:
    PdfReader = _get_pdf_reader_class()
    reader = PdfReader(BytesIO(data))
    if getattr(reader, "is_encrypted", False):
        raise RuntimeError("暂不支持加密 PDF，请先移除密码保护。")
    if len(reader.pages) > max_pdf_pages:
        raise ValueError(f"PDF 共 {len(reader.pages)} 页，超过上限 {max_pdf_pages} 页。")
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(f"--- 第 {index} 页 ---\n{text}")
    return "\n\n".join(pages).strip()


def _validate_docx_archive(
    data: bytes,
    *,
    max_uncompressed_bytes: int,
    max_compression_ratio: float,
) -> None:
    try:
        with zipfile.ZipFile(BytesIO(data)) as archive:
            members = archive.infolist()
            if len(members) > 10_000:
                raise ValueError("DOCX 压缩包文件项过多，已拒绝解析。")
            if any(item.flag_bits & 0x1 for item in members):
                raise ValueError("暂不支持包含加密文件项的 DOCX。")
            uncompressed = sum(item.file_size for item in members)
            compressed = sum(item.compress_size for item in members)
    except zipfile.BadZipFile as exc:
        raise ValueError("DOCX 文件不是有效的 Office 压缩包。") from exc

    if uncompressed > max_uncompressed_bytes:
        raise ValueError(
            f"DOCX 解压后约 {uncompressed / 1024 / 1024:.1f} MB，"
            f"超过上限 {max_uncompressed_bytes / 1024 / 1024:.1f} MB。"
        )
    ratio = uncompressed / max(compressed, 1)
    if ratio > max_compression_ratio:
        raise ValueError(
            f"DOCX 压缩比 {ratio:.1f}:1 异常，超过上限 {max_compression_ratio:.1f}:1。"
        )


def _extract_docx_text(
    data: bytes,
    *,
    max_uncompressed_bytes: int,
    max_compression_ratio: float,
) -> str:
    _validate_docx_archive(
        data,
        max_uncompressed_bytes=max_uncompressed_bytes,
        max_compression_ratio=max_compression_ratio,
    )
    document = Document(BytesIO(data))
    paragraphs = [
        paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()
    ]
    return "\n".join(paragraphs).strip()


def extract_text_from_bytes(
    filename: str,
    data: bytes,
    *,
    max_pdf_pages: int = DEFAULT_MAX_PDF_PAGES,
    max_docx_uncompressed_bytes: int = DEFAULT_MAX_DOCX_UNCOMPRESSED_BYTES,
    max_docx_compression_ratio: float = DEFAULT_MAX_DOCX_COMPRESSION_RATIO,
) -> str:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_DOCUMENT_EXTENSIONS:
        raise ValueError(f"暂不支持的文献格式：{extension or '未知格式'}")

    if extension == ".pdf":
        text = _extract_pdf_text(data, max_pdf_pages=max_pdf_pages)
    elif extension == ".docx":
        text = _extract_docx_text(
            data,
            max_uncompressed_bytes=max_docx_uncompressed_bytes,
            max_compression_ratio=max_docx_compression_ratio,
        )
    else:
        text = _decode_text_bytes(data)

    text = _normalize_text(text)
    if not text:
        raise RuntimeError(f"文献 {filename} 未提取到可用文本，请检查文件是否为扫描版或加密文档。")
    return text


def _truncate_text(text: str, max_chars: int) -> tuple[str, bool]:
    if len(text) <= max_chars:
        return text, False
    middle_marker = "\n\n[...文档中部采样...]\n\n"
    tail_marker = "\n\n[...文档结尾采样...]\n\n"
    marker_chars = len(middle_marker) + len(tail_marker)
    if max_chars <= marker_chars + 2:
        return text[:max_chars], True

    sample_budget = max_chars - marker_chars
    head_size = sample_budget // 3
    middle_size = sample_budget // 3
    tail_size = sample_budget - head_size - middle_size
    middle_start = max(0, (len(text) - middle_size) // 2)
    head = text[:head_size].rstrip()
    middle = text[middle_start : middle_start + middle_size].strip()
    tail = text[-tail_size:].lstrip()
    sampled = f"{head}{middle_marker}{middle}{tail_marker}{tail}"
    return sampled, True


def _parse_document_payloads(
    payloads: Iterable[tuple[str, bytes]],
    *,
    max_docs: int = DEFAULT_MAX_DOCS,
    max_chars_per_doc: int = DEFAULT_MAX_CHARS_PER_DOC,
    max_file_bytes: int = DEFAULT_MAX_FILE_BYTES,
    max_total_bytes: int = DEFAULT_MAX_TOTAL_BYTES,
    max_pdf_pages: int = DEFAULT_MAX_PDF_PAGES,
    max_docx_uncompressed_bytes: int = DEFAULT_MAX_DOCX_UNCOMPRESSED_BYTES,
    max_docx_compression_ratio: float = DEFAULT_MAX_DOCX_COMPRESSION_RATIO,
) -> list[ParsedDocument]:
    payload_list = list(payloads)
    if len(payload_list) > max_docs:
        raise ValueError(f"最多允许上传 {max_docs} 篇文献，当前为 {len(payload_list)} 篇。")

    total_bytes = sum(len(data) for _, data in payload_list)
    if total_bytes > max_total_bytes:
        raise ValueError(
            f"文献总大小约 {total_bytes / 1024 / 1024:.1f} MB，"
            f"超过上限 {max_total_bytes / 1024 / 1024:.1f} MB。"
        )
    parsed_documents: list[ParsedDocument] = []

    for name, data in payload_list:
        if len(data) > max_file_bytes:
            raise ValueError(
                f"文献 {name} 约 {len(data) / 1024 / 1024:.1f} MB，"
                f"超过单文件上限 {max_file_bytes / 1024 / 1024:.1f} MB。"
            )

        text = extract_text_from_bytes(
            name,
            data,
            max_pdf_pages=max_pdf_pages,
            max_docx_uncompressed_bytes=max_docx_uncompressed_bytes,
            max_docx_compression_ratio=max_docx_compression_ratio,
        )
        truncated_text, truncated = _truncate_text(text, max_chars_per_doc)
        parsed_documents.append(
            ParsedDocument(
                name=name,
                extension=Path(name).suffix.lower(),
                content=truncated_text,
                original_chars=len(text),
                extracted_chars=len(truncated_text),
                truncated=truncated,
            )
        )

    if not parsed_documents:
        raise ValueError("没有解析到任何文献内容，请至少上传一篇支持的文献文件。")

    return parsed_documents


def parse_uploaded_documents(
    uploaded_files: Sequence[UploadedFile],
    *,
    max_docs: int = DEFAULT_MAX_DOCS,
    max_chars_per_doc: int = DEFAULT_MAX_CHARS_PER_DOC,
    max_file_bytes: int = DEFAULT_MAX_FILE_BYTES,
    max_total_bytes: int = DEFAULT_MAX_TOTAL_BYTES,
    max_pdf_pages: int = DEFAULT_MAX_PDF_PAGES,
    max_docx_uncompressed_bytes: int = DEFAULT_MAX_DOCX_UNCOMPRESSED_BYTES,
    max_docx_compression_ratio: float = DEFAULT_MAX_DOCX_COMPRESSION_RATIO,
) -> list[ParsedDocument]:
    if len(uploaded_files) > max_docs:
        raise ValueError(f"最多允许上传 {max_docs} 篇文献，当前为 {len(uploaded_files)} 篇。")

    payloads: list[tuple[str, bytes]] = []
    observed_total = 0
    for uploaded_file in uploaded_files:
        name = getattr(uploaded_file, "name", "document")
        declared_size = getattr(uploaded_file, "size", None)
        if isinstance(declared_size, int):
            if declared_size > max_file_bytes:
                raise ValueError(
                    f"文献 {name} 约 {declared_size / 1024 / 1024:.1f} MB，"
                    f"超过单文件上限 {max_file_bytes / 1024 / 1024:.1f} MB。"
                )
            observed_total += declared_size
            if observed_total > max_total_bytes:
                raise ValueError("文献总大小超过允许上限。")
        data = uploaded_file.getvalue()
        payloads.append((name, data))

    return _parse_document_payloads(
        payloads,
        max_docs=max_docs,
        max_chars_per_doc=max_chars_per_doc,
        max_file_bytes=max_file_bytes,
        max_total_bytes=max_total_bytes,
        max_pdf_pages=max_pdf_pages,
        max_docx_uncompressed_bytes=max_docx_uncompressed_bytes,
        max_docx_compression_ratio=max_docx_compression_ratio,
    )


def parse_document_paths(
    file_paths: Sequence[str],
    *,
    max_docs: int = DEFAULT_MAX_DOCS,
    max_chars_per_doc: int = DEFAULT_MAX_CHARS_PER_DOC,
    max_file_bytes: int = DEFAULT_MAX_FILE_BYTES,
    max_total_bytes: int = DEFAULT_MAX_TOTAL_BYTES,
    max_pdf_pages: int = DEFAULT_MAX_PDF_PAGES,
    max_docx_uncompressed_bytes: int = DEFAULT_MAX_DOCX_UNCOMPRESSED_BYTES,
    max_docx_compression_ratio: float = DEFAULT_MAX_DOCX_COMPRESSION_RATIO,
) -> list[ParsedDocument]:
    if len(file_paths) > max_docs:
        raise ValueError(f"最多允许读取 {max_docs} 篇文献，当前为 {len(file_paths)} 篇。")

    payloads: list[tuple[str, bytes]] = []
    total_bytes = 0
    for raw_path in file_paths:
        path = Path(raw_path).expanduser()
        if not path.is_file():
            raise FileNotFoundError(f"文献文件不存在：{path}")
        size = path.stat().st_size
        if size > max_file_bytes:
            raise ValueError(f"文献 {path.name} 超过单文件大小上限。")
        total_bytes += size
        if total_bytes > max_total_bytes:
            raise ValueError("文献总大小超过允许上限。")
        payloads.append((path.name, path.read_bytes()))

    return _parse_document_payloads(
        payloads,
        max_docs=max_docs,
        max_chars_per_doc=max_chars_per_doc,
        max_file_bytes=max_file_bytes,
        max_total_bytes=max_total_bytes,
        max_pdf_pages=max_pdf_pages,
        max_docx_uncompressed_bytes=max_docx_uncompressed_bytes,
        max_docx_compression_ratio=max_docx_compression_ratio,
    )


def build_paper_preview(
    documents: Sequence[ParsedDocument],
    *,
    preview_chars: int = DEFAULT_PREVIEW_CHARS,
) -> str:
    sections: list[str] = []
    for index, document in enumerate(documents, start=1):
        preview = document.content[:preview_chars].rstrip()
        if len(document.content) > preview_chars:
            preview += "..."
        sections.append(
            f"""### 文献 {index}：{document.name}
- 类型：{document.extension}
- 提取长度：{document.original_chars} 字符

{preview}
"""
        )
    return "\n\n".join(sections).strip()
